import streamlit as st
import datetime
from docx import Document
import io

st.set_page_config(page_title="RheumaView AI Lite — Debug", layout="wide")

st.title("RheumaView AI Lite")
st.subheader("Radiologic Reasoning for Rheumatologists")
st.caption("Stable debug version with interface restored. AI temporarily disabled. No clinical context in report.")

st.markdown("**Curated by Dr Olga Goodman**")
st.markdown("---")

st.markdown("#### Step 1: Upload Current Imaging")
current_images = st.file_uploader(
    "Upload up to 5 current radiographic images (jpg/png/webp only)",
    type=["jpg", "jpeg", "png", "webp"],
    accept_multiple_files=True,
    key="current_images"
)

st.markdown("#### Step 2: Patient Demographics")
dob = st.text_input("Date of Birth (optional, YYYY-MM-DD):")
age = st.number_input("Patient Age:", min_value=0, max_value=130, step=1)
sex = st.radio("Sex at Birth:", ["Female", "Male", "Other / Intersex"])
name = st.text_input("Patient Name or ID (optional):")
mrn = st.text_input("Medical Record Number (optional):")
clinical_context = st.text_area("Clinical context (for reference only, not included in report)")

st.markdown("#### Step 3: Select Regions to Analyze")
region = st.multiselect("Select anatomical regions shown in the uploaded images:",
                        ["Hands", "Wrists", "Feet", "Ankles", "Knees", "Shoulders", "Pelvis/SI Joints", "Spine (C/L/T)", "Other"])
report_type = st.radio("Report type:", ["Single report", "Report with interval change analysis"])

st.markdown("#### Step 4: Report Formatting Options")
custom_header = st.checkbox("Include custom header and footer?")
header_text = st.text_input("Default Header (editable later):") if custom_header else ""
footer_text = st.text_input("Default Footer (editable later):") if custom_header else ""

confirm = st.checkbox("I confirm that all relevant imaging has been uploaded.")
if st.button("READY to generate report") and confirm:
    st.markdown("---")
    st.subheader("📝 Generated Report Preview")

    if current_images and len(current_images) > 5:
        st.warning("Please upload no more than 5 images at a time.")
        st.stop()

    try:
        if not current_images:
            st.warning("No images uploaded. Please upload radiographic images.")
        else:
            st.success(f"{len(current_images)} image(s) uploaded. (AI analysis disabled in this version)")
            for file in current_images:
                st.write(f"• {file.name}")

            doc = Document()
            doc.add_heading("Radiology Report", 0)
            doc.add_paragraph(f"Date of Study: {str(datetime.date.today())}")
            doc.add_paragraph(f"Patient Age: {age}, Sex: {sex}")
            doc.add_paragraph(f"Regions: {', '.join(region)}")
            doc.add_paragraph("\nReport Content:\n")
            doc.add_paragraph("This is a placeholder report. AI-based image interpretation is currently disabled in this debug version.")

            if custom_header:
                doc.add_paragraph(f"\nHeader: {header_text}")
                doc.add_paragraph(f"\nFooter: {footer_text}")

            buffer = io.BytesIO()
            doc.save(buffer)
            buffer.seek(0)

            st.download_button(
                label="📥 Download .docx Report",
                data=buffer,
                file_name="radiology_report_debug.docx",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
            )

    except Exception as e:
        st.error(f"⚠️ Report generation failed: {str(e)}")
        st.stop()