
import streamlit as st
import datetime
from docx import Document

st.set_page_config(page_title="RheumaView AI Lite — FAILPROOF", layout="wide")

st.title("RheumaView AI Lite — Failproof Mode")
st.caption("Findings explicitly compiled, previewed, and exported.")

# --- Demographics ---
study_date = st.date_input("Date of current study", datetime.date.today())
age = st.number_input("Patient Age", min_value=0, max_value=120, step=1)
sex = st.radio("Sex at Birth", ["Female", "Male", "Other / Intersex"])
name = st.text_input("Patient Name", value="Name Name")
mrn = st.text_input("Medical Record Number", value="1234567")
dob_input = st.text_input("Date of Birth (YYYY-MM-DD)", value="2000-01-01")
st.text_area("Clinical context (for interface only; not in report)", value="chronic worsening low back pain", key="clinical_context")

regions = st.multiselect("Select regions to include findings for", [
    "SI Joints", "Spine (DISH)", "Peripheral Joints"
], default=["SI Joints", "Spine (DISH)", "Peripheral Joints"])

st.markdown("---")

# --- Manual findings compiler ---
default_text = ""
if "SI Joints" in regions:
    default_text += "**SI Joints**\n[Mild/Moderate/Advanced sacroiliitis description here]\n\n"

if "Spine (DISH)" in regions:
    default_text += "**Spine (DISH)**\n[Flowing ossification, preserved disc spaces]\n\n"

if "Peripheral Joints" in regions:
    default_text += "**Peripheral Joints**\nJoint space narrowing: \nErosions: \nOsteophytes: \nImpression: \n\n"

findings = st.text_area("✍️ Final Report Findings (editable, included in DOCX):", value=default_text, height=300)

# --- Generate report ---
if st.button("Generate Final Word Report"):
    doc = Document()
    doc.add_paragraph("My Header")
    doc.add_heading("RheumaView Structured Report", 0)
    doc.add_paragraph(f"Date of Current Study: {study_date}")
    doc.add_paragraph(f"Patient: {name}")
    doc.add_paragraph(f"DOB: {dob_input}")
    doc.add_paragraph(f"MRN: {mrn}")
    doc.add_paragraph(f"Age: {age}")
    doc.add_paragraph(f"Sex: {sex}")
    doc.add_heading("RheumaView Report", level=1)
    doc.add_paragraph(findings.strip())
    doc.add_heading("Interval Comparison", level=2)
    doc.add_paragraph("Prior_1 (): Compared for progression/regression relative to current study.")
    doc.add_paragraph("\nMy Footer")

    filename = "rheumaview_structured_report_FAILPROOF.docx"
    doc.save(filename)

    with open(filename, "rb") as file:
        st.download_button("✅ Download DOCX Report", data=file, file_name=filename, mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document")

    st.success("✅ Report generated with visible, locked-in findings.")
