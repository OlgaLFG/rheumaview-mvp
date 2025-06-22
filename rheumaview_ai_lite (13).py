import streamlit as st
import datetime
from docx import Document

# --- CONFIGURATION ---
st.set_page_config(page_title="RheumaView AI Lite", layout="wide")
st.title("RheumaView AI Lite")
st.subheader("Radiologic Reasoning for Rheumatologists")
st.caption("Crash-protected interface. AI analysis currently disabled.")

# --- Step 1: Upload Current Imaging ---
st.markdown("### Step 1: Upload Current Imaging")
uploaded_images = st.file_uploader(
    "Upload current radiographic images (multiple files allowed)",
    type=["jpg", "jpeg", "png", "webp", "dcm", "bmp", "tiff"],
    accept_multiple_files=True,
    key="current_images"
)

# --- Step 2: Patient Info ---
st.markdown("### Step 2: Patient Demographics")
dob = st.text_input("Date of Birth (optional, YYYY-MM-DD):", value="")
age_input = st.text_input("Patient Age (in years):", value="")
sex = st.selectbox("Sex at Birth:", ["Female", "Male", "Other"])
name = st.text_input("Patient Name or ID (optional):", value="")
mrn = st.text_input("Medical Record Number (optional):", value="")
st.text_area("Clinical context (for reference only — not included in report)", value="", key="clinical_context")

# --- Step 3: Region and Type ---
st.markdown("### Step 3: Select Region and Report Type")
region_options = [
    "Multiple Regions", "Cervical Spine", "Thoracic Spine", "Lumbar Spine", "Pelvis/SI/Sacrum",
    "Hips", "Knees", "Ankles", "Feet", "Hands", "Wrists", "Elbows", "Shoulders", "Other"
]
region = st.selectbox("Region(s) to be analyzed:", region_options)

report_type = st.selectbox("Report Type:", ["Single Report (default)", "Interval Comparison (disabled)"])
if report_type == "Interval Comparison (disabled)":
    st.info("This version does not support prior imaging comparison.")

# --- Step 4: Header/Footer ---
st.markdown("### Step 4: Header/Footer (optional)")
custom_header = st.text_input("Custom Header:", value="RheumaView")
custom_footer = st.text_input("Custom Footer:", value="Curated by Dr. Olga Goodman")

# --- Step 5: Confirmation & Report ---
st.markdown("### Final Step: Confirm and Generate Report")
user_ready = st.checkbox("READY")

if st.button("Generate Report"):
    if not user_ready:
        st.warning("Please confirm readiness by checking the READY box.")
    elif not uploaded_images:
        st.warning("Please upload at least one image before generating the report.")
    else:
        try:
            age = int(age_input) if age_input.isdigit() else "N/A"

            doc = Document()
            doc.add_heading("RheumaView Structured Radiology Report", 0)
            doc.add_paragraph(f"Patient: {name if name else 'N/A'}")
            doc.add_paragraph(f"Age: {age}     Sex: {sex}")
            if dob:
                doc.add_paragraph(f"Date of Birth: {dob}")
            if mrn:
                doc.add_paragraph(f"MRN: {mrn}")
            doc.add_paragraph(f"Region analyzed: {region}")
            doc.add_paragraph(" ")
            doc.add_paragraph("Findings:")
            doc.add_paragraph("This is a placeholder report body. AI-based analysis is currently disabled.")
            doc.add_paragraph(" ")
            doc.add_paragraph("---")
            doc.add_paragraph(custom_header)
            doc.add_paragraph(custom_footer)

            output_path = "/mnt/data/rheumaview_structured_report_final.docx"
            doc.save(output_path)

            st.success("✅ Report generated successfully.")
            with open(output_path, "rb") as f:
                st.download_button("Download Report", f, file_name="rheumaview_structured_report.docx")

        except Exception as e:
            st.error("❌ Report generation failed. Please try again.")