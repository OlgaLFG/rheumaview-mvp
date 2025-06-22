import streamlit as st
from docx import Document
from datetime import datetime

st.set_page_config(page_title="RheumaView AI Lite", layout="wide")
st.title("RheumaView AI Lite")
st.subheader("Radiologic Reasoning for Rheumatologists")
st.markdown("⚠️ *AI interpretation is currently disabled. No identifiable data is stored. This version is crash-protected.*")

# --- Step 1: Upload Imaging ---
uploaded_images = st.file_uploader("Step 1: Upload current imaging", type=["jpg", "jpeg", "png", "webp", "bmp", "tiff", "dcm"], accept_multiple_files=True)

# --- Step 2: Demographics ---
st.markdown("### Step 2: Patient Demographics")
dob_str = st.text_input("Date of Birth (YYYY-MM-DD)")
age_input = st.text_input("Patient Age (optional, auto if DOB given)")
sex = st.selectbox("Sex at Birth", ["", "Female", "Male", "Other"])
name = st.text_input("Patient Name or ID (optional)")
mrn = st.text_input("Medical Record Number (optional)")
clinical_context = st.text_area("Clinical context (optional, not included in report)")

# --- Step 3: Region ---
st.markdown("### Step 3: Region(s) Shown")
region = st.selectbox("Anatomical region(s)", [
    "Multiple Regions", "Cervical Spine", "Thoracic Spine", "Lumbar Spine",
    "Pelvis/SI/Sacrum", "Hips", "Knees", "Feet", "Hands", "Shoulders"
])

# --- Step 4: Prior Imaging (interval analysis) ---
st.markdown("### Step 4: Comparison (optional)")
compare = st.radio("Compare with prior study?", ["No", "Yes"])
prior_date = ""
if compare == "Yes":
    prior_date = st.text_input("Enter prior study date (YYYY-MM-DD)")
    prior_images = st.file_uploader("Upload prior imaging files", type=["jpg", "jpeg", "png", "webp", "bmp", "tiff", "dcm"], accept_multiple_files=True, key="prior_upload")

# --- Step 5: Header/Footer ---
st.markdown("### Step 5: Header/Footer")
header = st.text_input("Header (editable)", value="")
footer = st.text_input("Footer (editable)", value="")

# --- Step 6: Confirmation ---
st.markdown("### Step 6: Confirm readiness to generate report")
ready = st.checkbox("READY – I confirm all data is entered")

# --- Generate Report ---
if st.button("Generate Report"):
    if not ready:
        st.warning("Please check the READY box before proceeding.")
    elif not uploaded_images:
        st.warning("Please upload current imaging files.")
    else:
        try:
            doc = Document()
            doc.add_heading("RheumaView Structured Radiology Report", 0)
            doc.add_paragraph(f"Patient: {name or 'N/A'}")
            doc.add_paragraph(f"Sex at Birth: {sex or 'N/A'}")

            # Process age
            age_final = "N/A"
            if dob_str:
                try:
                    dob = datetime.strptime(dob_str, "%Y-%m-%d")
                    today = datetime.today()
                    age_final = today.year - dob.year - ((today.month, today.day) < (dob.month, dob.day))
                except:
                    age_final = age_input or "N/A"
            else:
                age_final = age_input or "N/A"

            doc.add_paragraph(f"Age: {age_final}")
            if dob_str:
                doc.add_paragraph(f"Date of Birth: {dob_str}")
            if mrn:
                doc.add_paragraph(f"MRN: {mrn}")
            doc.add_paragraph(f"Region Analyzed: {region}")
            if compare == "Yes":
                doc.add_paragraph(f"Compared to prior imaging dated: {prior_date or 'Unknown'}")
            doc.add_paragraph("---")
            doc.add_paragraph("Findings:")
            doc.add_paragraph("This is a placeholder body. No AI interpretation is active in this version.")
            doc.add_paragraph("---")
            if header:
                doc.add_paragraph(header)
            if footer:
                doc.add_paragraph(footer)

            filename = "/mnt/data/rheumaview_structured_report_final.docx"
            doc.save(filename)
            st.success("✅ Report generated.")
            with open(filename, "rb") as f:
                st.download_button("Download Report", f, file_name="rheumaview_structured_report.docx")

        except Exception as e:
            st.error("An error occurred while generating the report.")