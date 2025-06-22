
import streamlit as st
import datetime
from docx import Document

st.set_page_config(page_title="RheumaView AI Lite — Final FIXED", layout="wide")

st.title("RheumaView AI Lite — Final FIXED")
st.caption("Stable autofill with verified export and no clinical context in report")

if "generate_mode" not in st.session_state:
    st.session_state.generate_mode = False

# --- Inputs ---
study_date = st.date_input("Date of current study", datetime.date.today())
age = st.number_input("Patient Age", min_value=0, max_value=120, step=1)
sex = st.radio("Sex at Birth", ["Female", "Male", "Other / Intersex"])
name = st.text_input("Patient Name", value="Name Name")
mrn = st.text_input("Medical Record Number", value="1234567")
dob_input = st.text_input("Date of Birth (YYYY-MM-DD)", value="2000-01-01")
st.text_area("Clinical context (for reference only, not included in report)", value="chronic worsening low back pain", key="clinical_context")

regions = st.multiselect("Select regions to analyze", [
    "Multiple Regions", "SI Joints", "Spine (DISH)", "Peripheral Joints"
], default=["Multiple Regions"])

if "Multiple Regions" in regions:
    expanded_regions = ["SI Joints", "Spine (DISH)", "Peripheral Joints"]
else:
    expanded_regions = regions

st.markdown("---")

if "SI Joints" in expanded_regions:
    st.radio("Confidence level: SI Joints", ["Low", "Moderate", "High"], key="si_level")

if "Spine (DISH)" in expanded_regions:
    st.radio("Confidence level: DISH", ["Moderate", "High"], key="dish_level")

if "Peripheral Joints" in expanded_regions:
    st.text_input("Joint space narrowing (symmetry + location):", key="joint_space")
    st.text_input("Erosions (present/absent, location, character):", key="erosions")
    st.selectbox("Bone density:", ["Normal", "Juxta-articular osteopenia"], key="density")
    st.selectbox("Periosteal reaction:", ["Absent", "Present"], key="periosteal")
    st.text_input("Soft tissue (swelling, masses, calcifications):", key="soft_tissue")
    st.selectbox("Osteophytes:", ["None", "Mild", "Moderate", "Severe"], key="osteophytes")
    st.text_input("Other findings (ankylosis, subluxation, etc.):", key="other")
    st.text_area("Impression:", height=150, key="impression")

if st.button("Generate Report"):
    st.session_state.generate_mode = True

if st.session_state.generate_mode:
    compiled = ""

    if "SI Joints" in expanded_regions:
        lvl = st.session_state.si_level
        if lvl == "Low":
            txt = "Sacroiliac joints are symmetric. Mild subchondral sclerosis without erosions or joint space narrowing. Findings may represent early sacroiliitis or degenerative change."
        elif lvl == "Moderate":
            txt = "Mild irregularity and sclerosis of bilateral sacroiliac joints, more pronounced on iliac sides. No erosions or ankylosis. Findings suggest early sacroiliitis."
        else:
            txt = "Bilateral sacroiliac joint sclerosis, erosions, and joint space narrowing. Findings consistent with Grade II–III sacroiliitis."
        compiled += "\n**SI Joints**\n" + txt + "\n"

    if "Spine (DISH)" in expanded_regions:
        lvl = st.session_state.dish_level
        if lvl == "Moderate":
            txt = "Anterior vertebral body osteophytes and subtle ossification along the anterior longitudinal ligament at T12–L2, with preserved disc spaces. Findings may be consistent with early DISH."
        else:
            txt = "Flowing ossification along the anterior longitudinal ligament spanning T12 to L3 with preserved disc heights and absence of significant degenerative change. Findings consistent with diffuse idiopathic skeletal hyperostosis (DISH)."
        compiled += "\n**Spine (DISH)**\n" + txt + "\n"

    if "Peripheral Joints" in expanded_regions:
        compiled += f"""\n**Peripheral Joints**
Joint space narrowing: {st.session_state.joint_space}
Erosions: {st.session_state.erosions}
Bone density: {st.session_state.density}
Periosteal reaction: {st.session_state.periosteal}
Soft tissue: {st.session_state.soft_tissue}
Osteophytes: {st.session_state.osteophytes}
Other: {st.session_state.other}
Impression: {st.session_state.impression}\n"""

    st.markdown("### ✅ Preview of Findings")
    st.code(compiled.strip())

    doc = Document()
    doc.add_paragraph("My Header")
    doc.add_heading("RheumaView Structured Report", 0)
    doc.add_paragraph(f"Date of Current Study: {study_date}")
    doc.add_paragraph(f"Patient: {name}")
    doc.add_paragraph(f"DOB: {dob_input}")
    doc.add_paragraph(f"MRN: {mrn}")
    doc.add_paragraph(f"Age: {age}")
    doc.add_paragraph(f"Sex: {sex}")
    doc.add_heading("RheumaView Report", level=1)
    doc.add_paragraph(compiled.strip())
    doc.add_heading("Interval Comparison", level=2)
    doc.add_paragraph("Prior_1 (2025-06-18): Compared for progression/regression relative to current study.")
    doc.add_paragraph("\nMy Footer")

    filename = "rheumaview_structured_report_FINAL_FIXED.docx"
    doc.save(filename)
    with open(filename, "rb") as f:
        st.download_button("Download Word Report ✅", data=f, file_name=filename, mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document")

    st.success("✅ Report successfully compiled.")
