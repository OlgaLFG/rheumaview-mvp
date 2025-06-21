
import streamlit as st
import datetime
from docx import Document

# --- PAGE CONFIG ---
st.set_page_config(page_title="RheumaView AI Lite (Auto)", layout="wide")

# --- HEADER ---
st.title("RheumaView AI Lite")
st.subheader("Auto-filled Structured Report with AI-style Templates")
st.caption("Curated by Dr. Olga Goodman")

# --- STUDY DATA ---
today_default = datetime.date.today()
study_date = st.date_input("Date of current study:", value=today_default)
age = st.number_input("Patient Age:", min_value=0, max_value=120, step=1, value=50)
sex = st.radio("Sex at Birth:", ["Female", "Male", "Other / Intersex"])
name = st.text_input("Patient Name:")
mrn = st.text_input("Medical Record Number:")
dob_input = st.text_input("Date of Birth (YYYY-MM-DD):", value="2000-01-01")
clinical_context = st.text_area("Clinical context:", height=100)

regions = st.multiselect("Select regions to analyze:", [
    "Multiple Regions", "SI Joints", "Spine (DISH)", "Peripheral Joints"
], default=["Multiple Regions"])

st.markdown("---")

# --- Expand logic ---
if "Multiple Regions" in regions:
    expanded_regions = ["SI Joints", "Spine (DISH)", "Peripheral Joints"]
else:
    expanded_regions = regions

compiled_findings = ""

# --- AI SUGGESTION BLOCKS ---

if "SI Joints" in expanded_regions:
    st.header("Sacroiliac Joints (SI)")
    si_level = st.radio("Confidence level (SI):", ["Low", "Moderate", "High"], key="si_level")
    if si_level == "Low":
        text = ("Sacroiliac joints are symmetric. Mild subchondral sclerosis without erosions or joint space narrowing. "
                "Findings may represent early sacroiliitis or degenerative change.")
    elif si_level == "Moderate":
        text = ("Mild irregularity and sclerosis of bilateral sacroiliac joints, more pronounced on iliac sides. "
                "No erosions or ankylosis. Findings suggest early sacroiliitis.")
    else:
        text = ("Bilateral sacroiliac joint sclerosis, erosions, and joint space narrowing. "
                "Findings consistent with Grade II–III sacroiliitis.")
    st.text_area("AI-style phrasing:", value=text, height=150, key="si_text")
    compiled_findings += "\n**SI Joints**\n" + text + "\n"

if "Spine (DISH)" in expanded_regions:
    st.header("Spinal Findings (DISH)")
    dish_level = st.radio("Confidence level (DISH):", ["Moderate", "High"], key="dish_level")
    if dish_level == "Moderate":
        text = ("Anterior vertebral body osteophytes and subtle ossification along the anterior longitudinal ligament "
                "at T12–L2, with preserved disc spaces. Findings may be consistent with early DISH.")
    else:
        text = ("Flowing ossification along the anterior longitudinal ligament spanning T12 to L3 with preserved disc "
                "heights and absence of significant degenerative change. Findings consistent with diffuse idiopathic "
                "skeletal hyperostosis (DISH).")
    st.text_area("AI-style phrasing:", value=text, height=150, key="dish_text")
    compiled_findings += "\n**Spine (DISH)**\n" + text + "\n"

if "Peripheral Joints" in expanded_regions:
    st.header("Peripheral Joints")
    st.markdown("##### SUGGESTED REPORTING TEMPLATE (Plain Radiograph)")
    joint_space = st.text_input("Joint space narrowing (symmetry + location):")
    erosions = st.text_input("Erosions (present/absent, location, character):")
    density = st.selectbox("Bone density:", ["Normal", "Juxta-articular osteopenia"])
    periosteal = st.selectbox("Periosteal reaction:", ["Absent", "Present"])
    soft_tissue = st.text_input("Soft tissue (swelling, masses, calcifications):")
    osteophytes = st.selectbox("Osteophytes:", ["None", "Mild", "Moderate", "Severe"])
    other = st.text_input("Other findings (ankylosis, subluxation, etc.):")
    impression = st.text_area("Impression:",
        "Imaging features are [ / are not ] consistent with an inflammatory arthropathy.\n"
        "The findings may suggest [RA / PsA / gout / EOA / OA], based on [key features].", height=150)

    compiled_findings += f"""\n**Peripheral Joints**
Joint space narrowing: {joint_space}
Erosions: {erosions}
Bone density: {density}
Periosteal reaction: {periosteal}
Soft tissue: {soft_tissue}
Osteophytes: {osteophytes}
Other: {other}
Impression: {impression}\n"""

st.markdown("---")
confirmed = st.checkbox("Confirm all data is ready for report generation")

if confirmed and st.button("Generate DOCX Report"):
    doc = Document()
    doc.add_paragraph("my header")
    doc.add_heading("RheumaView Structured Report", 0)
    doc.add_paragraph(f"Date of Current Study: {study_date}")
    doc.add_paragraph(f"Patient: {name}")
    doc.add_paragraph(f"DOB: {dob_input}")
    doc.add_paragraph(f"MRN: {mrn}")
    doc.add_paragraph(f"Age: {age}")
    doc.add_paragraph(f"Sex: {sex}")
    doc.add_paragraph(f"Clinical context: {clinical_context}")
    doc.add_heading("RheumaView Report", level=1)
    doc.add_paragraph(compiled_findings)

    doc.add_heading("Interval Comparison", level=2)
    doc.add_paragraph("Prior_1 (2025-06-19): Compared for progression/regression relative to current study.")

    doc.add_paragraph("\nmy footer")
    filename = "rheumaview_structured_autofilled.docx"
    doc.save(filename)

    with open(filename, "rb") as file:
        st.download_button(
            label="Download Word Report",
            data=file,
            file_name=filename,
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        )
