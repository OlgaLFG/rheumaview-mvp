
import streamlit as st
import datetime
from docx import Document

# --- CONFIGURATION ---
st.set_page_config(page_title="RheumaView AI Lite", layout="wide")
st.title("RheumaView AI Lite")
st.subheader("Radiologic Reasoning for Rheumatologists")
st.caption("Stable UI version with safe findings export, 'Multiple Regions' support, and interval analysis placeholder.")

# --- Step 1: Upload Current Imaging ---
st.markdown("### Step 1: Upload Current Imaging")
uploaded_images = st.file_uploader(
    "Upload current radiographic images (multiple files allowed)",
    type=["jpg", "jpeg", "png", "dcm", "webp", "bmp", "tiff"],
    accept_multiple_files=True,
    key="current_images"
)

# --- Step 2: Demographic Info ---
st.markdown("### Step 2: Patient Info")
patient_name = st.text_input("Patient Name or Initials:")
patient_age = st.number_input("Age:", min_value=0, max_value=120, step=1)
patient_sex = st.selectbox("Sex at birth:", ["Male", "Female", "Other"])
mrn = st.text_input("Medical Record Number (optional):")
clinical_context = st.text_area("Clinical context (for reference only, not included in report)")

# --- Step 3: Report Logic ---
st.markdown("### Step 3: Select Regions to Analyze")
region_options = [
    "Multiple Regions",
    "Cervical Spine", "Thoracic Spine", "Lumbar Spine", "Pelvis/SI/Sacrum",
    "Hips", "Knees", "Ankles", "Feet",
    "Hands", "Wrists", "Elbows", "Shoulders",
    "Other"
]
selected_region = st.selectbox("Select anatomical regions shown in the uploaded images:", region_options)

report_type = st.radio("Report type:", ["Single report", "Report with interval change analysis"])
if report_type == "Report with interval change analysis":
    st.warning("Interval comparison currently disabled in this version.")
    st.stop()

# --- Step 4: Formatting ---
st.markdown("### Step 4: Report Formatting Options")
use_custom_header = st.checkbox("Include custom header and footer?")

# --- Step 5: Generate Report ---
if st.button("Generate Report"):
    if not uploaded_images:
        st.error("Please upload at least one image.")
        st.stop()

    doc = Document()
    doc.add_heading("RheumaView Structured Radiology Report", 0)
    doc.add_paragraph(f"Patient: {patient_name}    Age: {patient_age}    Sex: {patient_sex}")
    if mrn:
        doc.add_paragraph(f"MRN: {mrn}")
    doc.add_paragraph(f"Region analyzed: {selected_region}")
    doc.add_paragraph(" ")
    doc.add_paragraph("Findings:")
    doc.add_paragraph("This is a placeholder report body. AI analysis is currently disabled.")
    doc.add_paragraph(" ")

    if use_custom_header:
        doc.add_paragraph("----- Custom Header/Footer Enabled -----")

    report_path = "/mnt/data/rheumaview_structured_report_output.docx"
    doc.save(report_path)
    st.success("Report generated successfully.")
    with open(report_path, "rb") as f:
        st.download_button("Download Report", f, file_name="rheumaview_structured_report.docx")
