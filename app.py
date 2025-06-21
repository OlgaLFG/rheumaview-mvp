
import streamlit as st
from fpdf import FPDF
from PIL import Image
import datetime

# --- PAGE CONFIG ---
st.set_page_config(page_title="RheumaView", page_icon=":green_square:", layout="wide")

# --- LOGO ---
logo = Image.open("logo.png")
st.image(logo, width=100)

# --- HEADER ---
st.title("🟩 RheumaView™")
st.subheader("Radiologic Reasoning for Rheumatologists")
st.caption("Curated by Dr. Olga Goodman")

st.markdown("### Step 1: Upload Current Imaging")
uploaded_current = st.file_uploader(
    "Upload current radiographic images (multiple files allowed)", 
    type=["jpg", "jpeg", "png", "dcm"], 
    accept_multiple_files=True
)

today_default = datetime.date.today()
study_date = st.date_input("Date of current study:", value=today_default)

st.markdown("### Step 2: Select Regions to Analyze")
regions = st.multiselect("Select all anatomical regions shown in the uploaded images:", [
    "Cervical Spine", "Thoracic Spine", "Lumbar Spine", "Pelvis / SI joints",
    "Hip", "Knee", "Ankle", "Foot", "Hand", "Wrist", "Elbow", "Shoulder"
])

mode = st.radio("Report type:", ["Single report", "Report with interval change analysis"])

compare_enabled = False
prior_images = {}
if mode == "Report with interval change analysis":
    compare_enabled = st.checkbox("Compare with prior imaging?")
    if compare_enabled:
        st.markdown("### Step 3: Upload Prior Imaging for Comparison")
        num_priors = st.number_input("How many prior studies to upload?", min_value=1, max_value=5, step=1)
        for i in range(num_priors):
            st.markdown(f"**Prior Study {i+1}:**")
            files = st.file_uploader(
                f"Upload image files for prior study #{i+1}", 
                type=["jpg", "jpeg", "png", "dcm"], 
                accept_multiple_files=True, 
                key=f"prior_{i}"
            )
            prior_date = st.text_input(
                f"Enter date of prior study #{i+1} (full date or year):", 
                value="", 
                key=f"date_{i}"
            )
            if files:
                prior_images[f"Prior_{i+1}"] = {"files": files, "date": prior_date}

confirmed_all = st.checkbox("I confirm that all relevant imaging has been uploaded.")

st.markdown("---")
if confirmed_all:
    ready = st.button("✅ READY to generate report")
else:
    st.warning("Please confirm that all imaging has been uploaded before proceeding.")

if confirmed_all and ready:
    st.success("📝 Generating structured report...")

    findings = {}
    for region in regions:
        st.markdown(f"#### {region}")
        findings[region] = st.text_area(f"Enter detailed findings for {region}:", height=150)

    # Summary option
    summary = st.text_area("Optional: Add a brief EMR-friendly summary (will be included separately):", height=200)

    # Generate report text
    full_report = f"RheumaView™ Report\nDate of Current Study: {study_date}\n\n"
    for region, text in findings.items():
        full_report += f"---\nRegion: {region}\n{text}\n"

    if compare_enabled and prior_images:
        full_report += "\n---\nInterval Comparison:\n"
        for label, data in prior_images.items():
            full_report += f"{label} ({data['date']}): Compared for progression/regression relative to current study.\n"

    # Append summary
    import streamlit as st
from docx import Document
from PIL import Image
import datetime
import os

# --- PAGE CONFIG ---
st.set_page_config(page_title="RheumaView", page_icon=None, layout="wide")

# --- DISCLAIMER ---
st.warning("""
⚠️ **Disclaimer:**
This app does not perform automated radiologic interpretation or image recognition. Any regions listed below are for interface structure only.
The text content of this report is manually entered by the user.
Any subsequent edits made after download are not monitored or controlled by the system.
No identifiable patient data is stored on this server.
All patient-specific information entered here is processed locally and only saved to user device upon download.
""")

# --- LOGO ---
logo = Image.open("logo.png")
st.image(logo, width=100)

# --- HEADER ---
st.title("RheumaView")
st.subheader("Radiologic Reasoning for Rheumatologists")
st.caption("Curated by Dr. Olga Goodman")

st.markdown("### Step 1: Upload Current Imaging")
uploaded_current = st.file_uploader(
    "Upload current radiographic images (multiple files allowed)", 
    type=["jpg", "jpeg", "png", "dcm", "webp", "bmp", "tiff"], 
    accept_multiple_files=True
)

today_default = datetime.date.today()
study_date = st.date_input("Date of current study:", value=today_default)

st.markdown("### Step 2: Patient Demographics")
dob = st.text_input("Date of Birth (optional, YYYY-MM-DD):")
calculated_age = ""
if dob:
    try:
        birth_date = datetime.datetime.strptime(dob, "%Y-%m-%d").date()
        today = datetime.date.today()
        calculated_age = today.year - birth_date.year - ((today.month, today.day) < (birth_date.month, birth_date.day))
    except:
        calculated_age = ""

age = st.number_input("Patient Age:", min_value=0, max_value=120, step=1, value=calculated_age if isinstance(calculated_age, int) else 0)
sex = st.radio("Sex at Birth:", ["Female", "Male", "Other / Intersex"])
name = st.text_input("Patient Name or ID (optional):")
mrn = st.text_input("Medical Record Number (optional):")
clinical_context = st.text_area("Optional: Clinical summary (symptoms, known diagnoses, exam findings, etc.)",
                                max_chars=10000, height=150)

st.markdown("### Step 3: Select Regions to Analyze")
region_options = [
    "Multiple Regions",
    "Cervical Spine", "Thoracic Spine", "Lumbar Spine", "Pelvis / SI joints",
    "Hip", "Knee", "Ankle", "Foot", "Hand", "Wrist", "Elbow", "Shoulder",
    "Other Regions (e.g., ribs, clavicle, forearm, chest, etc.)"
]
regions = st.multiselect("Select anatomical regions shown in the uploaded images:", region_options, default=["Multiple Regions"])

mode = st.radio("Report type:", ["Single report", "Report with interval change analysis"])

compare_enabled = False
prior_images = {}
if mode == "Report with interval change analysis":
    compare_enabled = st.checkbox("Compare with prior imaging?")
    if compare_enabled:
        st.markdown("### Step 4: Upload Prior Imaging for Comparison")
        num_priors = st.number_input("How many prior studies to upload?", min_value=1, max_value=5, step=1)
        for i in range(num_priors):
            st.markdown(f"**Prior Study {i+1}:**")
            files = st.file_uploader(
                f"Upload image files for prior study #{i+1}", 
                type=["jpg", "jpeg", "png", "dcm", "webp", "bmp", "tiff"], 
                accept_multiple_files=True, 
                key=f"prior_{i}"
            )
            prior_date = st.text_input(
                f"Enter date of prior study #{i+1} (full date or year):", 
                value="", 
                key=f"date_{i}"
            )
            if files:
                prior_images[f"Prior_{i+1}"] = {"files": files, "date": prior_date}

st.markdown("### Step 4: Report Formatting Options")
include_header_footer = st.checkbox("Include custom header and footer?")

import streamlit as st
from docx import Document
from PIL import Image
import datetime
import os

# --- PAGE CONFIG ---
st.set_page_config(page_title="RheumaView", page_icon=None, layout="wide")

# --- DISCLAIMER ---
st.warning("""
⚠️ **Disclaimer:**
This app does not perform automated radiologic interpretation or image recognition. Any regions listed below are for interface structure only.
The text content of this report is manually entered by the user.
Any subsequent edits made after download are not monitored or controlled by the system.
No identifiable patient data is stored on this server.
All patient-specific information entered here is processed locally and only saved to user device upon download.
""")

# --- LOGO ---
logo = Image.open("logo.png")
st.image(logo, width=100)

# --- HEADER ---
st.title("RheumaView")
st.subheader("Radiologic Reasoning for Rheumatologists")
st.caption("Curated by Dr. Olga Goodman")

st.markdown("### Step 1: Upload Current Imaging")
uploaded_current = st.file_uploader(
    "Upload current radiographic images (multiple files allowed)", 
    type=["jpg", "jpeg", "png", "dcm", "webp", "bmp", "tiff"], 
    accept_multiple_files=True
)

today_default = datetime.date.today()
study_date = st.date_input("Date of current study:", value=today_default)

st.markdown("### Step 2: Patient Demographics")
dob_input = st.text_input("Date of Birth (optional, YYYY-MM-DD):")
calculated_age = None
if dob_input:
    try:
        birth_date = datetime.datetime.strptime(dob_input, "%Y-%m-%d").date()
        today = datetime.date.today()
        calculated_age = today.year - birth_date.year - ((today.month, today.day) < (birth_date.month, birth_date.day))
    except ValueError:
        st.warning("Invalid date format. Use YYYY-MM-DD.")
        calculated_age = None

age = st.number_input("Patient Age:", min_value=0, max_value=120, step=1, value=calculated_age if calculated_age is not None else 0)
sex = st.radio("Sex at Birth:", ["Female", "Male", "Other / Intersex"])
name = st.text_input("Patient Name or ID (optional):")
mrn = st.text_input("Medical Record Number (optional):")
clinical_context = st.text_area("Optional: Clinical summary (symptoms, known diagnoses, exam findings, etc.)",
                                max_chars=10000, height=150)

st.markdown("### Step 3: Select Regions to Analyze")
region_options = [
    "Multiple Regions",
    "Cervical Spine", "Thoracic Spine", "Lumbar Spine", "Pelvis / SI joints",
    "Hip", "Knee", "Ankle", "Foot", "Hand", "Wrist", "Elbow", "Shoulder",
    "Other Regions (e.g., ribs, clavicle, forearm, chest, etc.)"
]
regions = st.multiselect("Select anatomical regions shown in the uploaded images:", region_options, default=["Multiple Regions"])

mode = st.radio("Report type:", ["Single report", "Report with interval change analysis"])

compare_enabled = False
prior_images = {}
if mode == "Report with interval change analysis":
    compare_enabled = st.checkbox("Compare with prior imaging?")
    if compare_enabled:
        st.markdown("### Step 4: Upload Prior Imaging for Comparison")
        num_priors = st.number_input("How many prior studies to upload?", min_value=1, max_value=5, step=1)
        for i in range(num_priors):
            st.markdown(f"**Prior Study {i+1}:**")
            files = st.file_uploader(
                f"Upload image files for prior study #{i+1}", 
                type=["jpg", "jpeg", "png", "dcm", "webp", "bmp", "tiff"], 
                accept_multiple_files=True, 
                key=f"prior_{i}"
            )
            prior_date = st.text_input(
                f"Enter date of prior study #{i+1} (full date or year):", 
                value="", 
                key=f"date_{i}"
            )
            if files:
                prior_images[f"Prior_{i+1}"] = {"files": files, "date": prior_date}

st.markdown("### Step 4: Report Formatting Options")
include_header_footer = st.checkbox("Include custom header and footer?")

default_header = st.text_input("Default Header (editable later):", value="")
default_footer = st.text_input("Default Footer (editable later):", value="")

if include_header_footer:
    custom_header = default_header
    custom_footer = default_footer
else:
    custom_header = ""
    custom_footer = ""

confirmed_all = st.checkbox("I confirm that all relevant imaging has been uploaded.")

st.markdown("---")
if confirmed_all:
    ready = st.button("READY to generate report")
else:
    st.warning("Please confirm that all imaging has been uploaded before proceeding.")

if confirmed_all and ready:
    st.success("Generating structured report...")

    allow_freeform = st.checkbox("Allow free-form description without selecting anatomical regions")

    findings = {}
    if not allow_freeform:
        for region in regions:
            st.markdown(f"#### {region}")
            findings[region] = st.text_area(f"Enter detailed findings for {region}:", height=150)
    else:
        freeform_text = st.text_area("Enter full findings (all regions or general impressions):", height=250)
        findings["General"] = freeform_text

    summary = st.text_area("Optional: Add a brief EMR-friendly summary (will be included separately):", height=200)

    # --- WORD DOCUMENT GENERATION ---
    doc = Document()

    if custom_header:
        doc.add_paragraph(custom_header)

    doc.add_heading("RheumaView Structured Report", 0)
    doc.add_paragraph(f"Date of Current Study: {study_date}")
    if name:
        doc.add_paragraph(f"Patient: {name}")
    if dob_input:
        doc.add_paragraph(f"DOB: {dob_input}")
    if mrn:
        doc.add_paragraph(f"MRN: {mrn}")
    doc.add_paragraph(f"Age: {age}")
    doc.add_paragraph(f"Sex: {sex}")

    if clinical_context:
        doc.add_paragraph(f"Clinical context: {clinical_context}")

    doc.add_heading("RheumaView Report", level=1)
    for region, text in findings.items():
        doc.add_heading(f"{region}", level=2)
        doc.add_paragraph(text)

    if compare_enabled and prior_images:
        doc.add_heading("Interval Comparison", level=2)
        for label, data in prior_images.items():
            doc.add_paragraph(f"{label} ({data['date']}): Compared for progression/regression relative to current study.")

    if summary.strip():
        doc.add_paragraph("")
        doc.add_paragraph("=== EMR Summary ===")
        doc.add_paragraph(summary.strip())

    if custom_footer:
        doc.add_paragraph("\n\n" + custom_footer)

    filename = f"rheumaview_structured_report_{study_date}_{sex}_{age}.docx"
    doc.save(filename)

    with open(filename, "rb") as file:
        st.download_button(
            label="Download Word Report",
            data=file,
            file_name=filename,
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        )
